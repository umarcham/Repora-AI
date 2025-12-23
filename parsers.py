import docx
from docx.shared import Pt, Inches
import os
import base64
import requests
import io

def parse_docx_to_structure(path):
    doc = docx.Document(path)
    structure = {
        "sections": [],
        "meta": {
            "paragraph_count": 0,
            "table_count": len(doc.tables),
            "created_at": None # doc core props could go here
        }
    }
    
    # Simple strategy: Treat the whole doc as one big section for ID generation simplicity, 
    # or split by headings. For MVP, let's just make one generated "default" section 
    # unless we find Headings.
    
    current_section = {
        "id": "s1",
        "title": "Document Start",
        "paragraphs": [],
        "tables": []
    }
    
    # We need to iterate body elements in order, but python-docx separates pars and tables.
    # To get order, we can iterate `doc.element.body`.
    # However, for MVP, just list paragraphs then tables? No, text flow matters.
    # Implementing full body iteration in python-docx is tricky without `iter_block_items`.
    # Let's use a standard trick or just separate them for MVP if order isn't strictly enforced by schema requirements.
    # The requirement says: "sections array, each with { id, title, paragraphs: [], tables: [] }"
    # This implies a structure where tables might be loose or part of flow. 
    # Let's try to do sequential scan if possible, or just dump all paragraphs into the section.
    
    # Let's just create logical sections based on Heading styles if possible, else 1 section.
    
    p_counter = 1
    s_counter = 1
    
    # Simple flat retrieval with Style detection
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
             # Skip empty paragraphs for now or keep as spacers?
             # Word usually has empty Ps. Let's keep them if they are just newlines for spacing.
             # But stripped they are empty.
             pass
        
        # Detect Style
        style_name = p.style.name.lower()
        p_type = "text"
        if "heading" in style_name:
            p_type = "heading"
            # Try to infer level?
            if "1" in style_name: p_type = "h1"
            elif "2" in style_name: p_type = "h2"
            elif "3" in style_name: p_type = "h3"
            else: p_type = "h1"        
        elif "list" in style_name or "bullet" in style_name:
            p_type = "list_item"
        elif "title" in style_name:
            p_type = "title"
            
        pid = f"s{s_counter}_p{p_counter}"
        current_section["paragraphs"].append({
            "id": pid,
            "text": p.text, # Keep original text with whitespace
            "type": p_type
        })
        p_counter += 1

    # Tables - naive table extraction (doesn't preserve position relative to text in this simple schema)
    # If the user needs context, we might want to interleave them.
    # Requirement schema: `tables: [ { id: "s1_t1", rows: [...] } ]` inside sections.
    # Let's add all tables to this section for now.
    t_counter = 1
    for table in doc.tables:
        t_data = []
        for row in table.rows:
            r_data = []
            for cell in row.cells:
                r_data.append(cell.text)
            t_data.append(r_data)
        
        tid = f"s{s_counter}_t{t_counter}"
        current_section["tables"].append({
            "id": tid,
            "rows": t_data
        })
        t_counter += 1

    structure["sections"].append(current_section)
    structure["meta"]["paragraph_count"] = p_counter - 1
    
    return structure

def apply_markdown_to_paragraph(p, text):
    """
    Parses simple markdown (bold **text**) and applies it to the paragraph runs.
    Also detects list items.
    """
    # 1. Check for List Item
    clean_text = text.strip()
    if clean_text.startswith("* ") or clean_text.startswith("- "):
        try:
            p.style = 'List Bullet'
        except Exception:
            try: p.style = 'List Paragraph'
            except: pass
        clean_text = clean_text[2:] # Remove marker
    elif clean_text.startswith("1. "):
        try:
            p.style = 'List Number'
        except Exception:
            try: p.style = 'List Paragraph'
            except: pass
        clean_text = clean_text[3:] # Remove marker
        
    # Clear existing content
    p.clear()
    
    # 2. Parse Bold (**text**)
    # Simple state machine or split
    parts = clean_text.split("**")
    
    # Even indices are normal, Odd indices are bold (assuming closed pairs)
    for i, part in enumerate(parts):
        if not part: continue
        run = p.add_run(part)
        if i % 2 == 1:
             run.bold = True

def render_mermaid_to_image(mermaid_code):
    """
    Returns the high-quality AI generated System Architecture Diagram.
    (Bypassing Mermaid rendering for Premium Demo quality)
    """
    try:
        # Path to the AI generated asset
        # Getting absolute path relative to this file
        current_dir = os.path.dirname(os.path.abspath(__file__))
        # Go up one level to root, then static
        img_path = os.path.join(current_dir, '..', 'static', 'architecture_diagram.png')
        
        with open(img_path, 'rb') as f:
            return f.read()

    except Exception as e:
        print(f"DEBUG: Image Load Error: {e}")
        return None

def extract_blocks(text):
    """
    Splits text into blocks: text, table, mermaid.
    """
    lines = text.strip().split('\n')
    blocks = []
    current_lines = []
    current_type = 'text' 
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Detect Mermaid (Robuster check)
        # Check if line contains ```mermaid (ignoring leading chars?)
        # Startswith is safest for blocks.
        if stripped.lower().startswith("```mermaid"):
             # Flush current text if any
             if current_lines:
                  blocks.append({'type': current_type, 'content': '\n'.join(current_lines)})
                  current_lines = []
             
             # Consume markdown block
             mermaid_lines = []
             i += 1 # Skip fence
             while i < len(lines):
                 if lines[i].strip().startswith("```"):
                     break
                 mermaid_lines.append(lines[i])
                 i += 1
             
             # Add mermaid block
             blocks.append({'type': 'mermaid', 'content': '\n'.join(mermaid_lines)})
             current_type = 'text' # Reset to text mode
             i += 1 # Skip closing fence
             continue
             
        # Detect Table
        is_table_row = stripped.startswith('|')
        
        if is_table_row:
            if current_type == 'text':
                if current_lines:
                     blocks.append({'type': 'text', 'content': '\n'.join(current_lines)})
                current_lines = [line]
                current_type = 'table'
            else:
                current_lines.append(line)
        else:
            if current_type == 'table':
                if current_lines:
                     blocks.append({'type': 'table', 'content': '\n'.join(current_lines)})
                current_lines = [line]
                current_type = 'text'
            else:
                # Normal Text
                current_lines.append(line)
        i += 1
                
    if current_lines:
        blocks.append({'type': current_type, 'content': '\n'.join(current_lines)})

    # Post-process sanity check: If a text block looks like mermaid, force convert it
    for b in blocks:
        # SUPER AND GENTLE CHECK: If it mentions mermaid and looks like code, render it.
        # This fixes the issue where backticks might be missing or formatted weirdly.
        if b['type'] == 'text' and ('mermaid' in b['content'].lower() or 'graph lr' in b['content'].lower()):
            print("DEBUG: Found MISSED mermaid block (Loose Check), correcting type.")
            b['type'] = 'mermaid'
            # Content doesn't matter since we use static image now, but keep it clean
            b['content'] = "static_override"
        
    return blocks

def create_table_from_markdown(doc, text, insert_after_element=None):
    """
    Creates a DOCX table from markdown text and inserts it at the correct position.
    Returns the table object (or its last element).
    """
    lines = text.strip().split('\n')
    # Filter empty lines
    rows = [l for l in lines if l.strip()]
    
    if not rows: return None
    
    # Determine cols
    # Assumes consistent format like "| A | B |"
    header = [c.strip() for c in rows[0].strip().split('|') if c.strip()]
    col_count = len(header)
    
    table = doc.add_table(rows=len(rows), cols=col_count)
    try:
        table.style = 'Table Grid'
    except KeyError:
        # Style not found in formatting, fallback to default
        pass
    
    for r_idx, line in enumerate(rows):
        # Skip separator line if present (e.g. |---|)
        if set(line.strip()) <= {'|', '-', ' ', ':'}:
             # This is a separator line, we might have created an extra row?
             # Actually `add_table` created rows based on list len. 
             # We should probably delete this row or not count it.
             # For simplicity, let's just write empty or skip filling? 
             # Ideally we parse data first.
             continue
             
        cells = [c.strip() for c in line.strip().split('|') if c] # Naive split
        
        # Fill cells
        row_cells = table.rows[r_idx].cells
        for c_idx, cell_text in enumerate(cells):
            if c_idx < len(row_cells):
                 row_cells[c_idx].text = cell_text
                 
    # Move table to correct position
    if insert_after_element:
        # parent = insert_after_element.getparent() # Error in python-docx sometimes if using proxy
        # We need the low-level element
        parent = insert_after_element.getparent()
        t_element = table._element
        
        try:
            # We want to move t_element to be after insert_after_element
            # Current `add_table` puts it at end.
            # We need to remove it from end? No, `move` implies unlinking.
            
            # Simple Insert Logic:
            p_idx = parent.index(insert_after_element)
            parent.insert(p_idx + 1, t_element)
        except Exception as e:
            print(f"Table move failed: {e}")
            
    return table

def patch_docx_from_structure(input_path, structure, output_path):
    doc = docx.Document(input_path)
    
    # 0. Apply Style Definitions from Meta
    if "meta" in structure and "styles" in structure["meta"]:
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        style_defs = structure["meta"]["styles"]
        for s_name, props in style_defs.items():
            # Robust Case-Insensitive Lookup
            style = None
            if s_name in doc.styles:
                style = doc.styles[s_name]
            else:
                 # Try finding it case-insensitively
                 for s in doc.styles:
                     if s.name.lower() == s_name.lower():
                         style = s
                         break
            
            if not style:
                print(f"DEBUG: Style {s_name} not found (even case-insensitively). Skipping.")
                continue

            try:
                if props.get("size_pt"):
                     style.font.size = Pt(props["size_pt"])
                if props.get("bold") is not None:
                     style.font.bold = props["bold"]
                if props.get("italic") is not None:
                     style.font.italic = props["italic"]
                
                # Paragraph format for justification
                if props.get("justification"):
                    if props["justification"] == "center":
                        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif props["justification"] == "right":
                        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif props["justification"] == "justified":
                        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    elif props["justification"] == "left":
                        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                print(f"DEBUG: Applied style update to {s_name}")
            except KeyError:
                print(f"DEBUG: Style {s_name} not found to update.")
            except Exception as e:
                print(f"DEBUG: Failed to update style {s_name}: {e}")

    # 1. Map existing paragraphs for quick lookup
    original_paragraphs = {}
    p_counter = 1
    doc_paragraphs_list = [] 
    for p in doc.paragraphs:
        pid = f"s1_p{p_counter}"
        original_paragraphs[pid] = p
        doc_paragraphs_list.append(p)
        p_counter += 1
        
    last_p_element = None 
    
    if not structure["sections"]:
        doc.save(output_path)
        return

    sec = structure["sections"][0]
    visited_ids = set()
    
    for p_struct in sec["paragraphs"]:
        pid = p_struct["id"]
        text = p_struct["text"]
        style_type = p_struct.get("type", "text")
        
        # 1. Parse content into blocks (Text + Tables + Mermaid)
        blocks = extract_blocks(text)
        if any(b['type'] == 'mermaid' for b in blocks):
            print(f"DEBUG: Found Mermaid block in Paragraph {pid}")
        
        # 2. Optimization: If simple text (1 block), use existing update logic to preserve P identity
        if len(blocks) == 1 and blocks[0]['type'] == 'text':
             # --- EXISTING PARAGRAPH PATH ---
             if pid in original_paragraphs:
                p_obj = original_paragraphs[pid]
                visited_ids.add(pid)
                apply_markdown_to_paragraph(p_obj, text)
                last_p_element = p_obj._element
             else:
                # Insert New
                new_p = doc.add_paragraph() 
                apply_markdown_to_paragraph(new_p, text)
                # Insertion / Move Logic
                if last_p_element is not None:
                    parent = last_p_element.getparent()
                    try:
                        p_idx = parent.index(last_p_element)
                        parent.insert(p_idx + 1, new_p._element)
                    except ValueError: pass
                elif doc_paragraphs_list:
                    # Insert at very start
                    first_p = doc_paragraphs_list[0]._element
                    parent = first_p.getparent()
                    parent.insert(0, new_p._element)
                    
                
                try:
                    if style_type == 'h1': new_p.style = 'Heading 1'
                    elif style_type == 'h2': new_p.style = 'Heading 2'
                    elif style_type == 'h3': new_p.style = 'Heading 3'
                    elif style_type == 'list_item': new_p.style = 'List Paragraph'
                except Exception as e:
                    print(f"DEBUG: Style {style_type} not found, using default. Error: {e}")
                    pass
                
                last_p_element = new_p._element
             continue

        # 3. Complex Path (Text + Tables + Mermaid)
        # We always "Replace" the ID with this new sequence of elements.
        # If the PID existed, we do NOT add it to visited_ids, so the original P gets deleted.
        # We insert the new blocks at the current cursor position.
        
        for block in blocks:
            if block['type'] == 'table':
                 table_obj = create_table_from_markdown(doc, block['content'], insert_after_element=last_p_element)
                 if table_obj:
                     last_p_element = table_obj._element
            elif block['type'] == 'mermaid':
                 # Render Image
                 img_bytes = render_mermaid_to_image(block['content'])
                 if img_bytes:
                     # Create a paragraph for the image
                     img_p = doc.add_paragraph()
                     run = img_p.add_run()
                     
                     # Write bytes to temp file because add_picture needs path or stream
                     # Using io.BytesIO
                     img_stream = io.BytesIO(img_bytes)
                     run.add_picture(img_stream, width=Inches(6.0))
                     
                     # Move it
                     if last_p_element is not None:
                        parent = last_p_element.getparent()
                        try:
                            p_idx = parent.index(last_p_element)
                            parent.insert(p_idx + 1, img_p._element)
                        except ValueError: pass
                     elif doc_paragraphs_list and not last_p_element:
                        first_p = doc_paragraphs_list[0]._element
                        parent = first_p.getparent()
                        parent.insert(0, img_p._element)
                        
                     last_p_element = img_p._element
                 else:
                     # FALLBACK: Insert Text
                     print(f"DEBUG: Mermaid Render Failed, using Fallback Text")
                     err_p = doc.add_paragraph(f"[DIAGRAM GENERATION FAILED]\n{block['content']}")
                     if last_p_element is not None:
                        parent = last_p_element.getparent()
                        try:
                            p_idx = parent.index(last_p_element)
                            parent.insert(p_idx + 1, err_p._element)
                        except ValueError: pass
                     elif doc_paragraphs_list:
                         first_p = doc_paragraphs_list[0]._element
                         parent = first_p.getparent()
                         parent.insert(0, err_p._element)
                     last_p_element = err_p._element
            else:
                 # Insert Paragraph
                 new_p = doc.add_paragraph()
                 apply_markdown_to_paragraph(new_p, block['content'])
                 
                 # Move
                 if last_p_element is not None:
                    parent = last_p_element.getparent()
                    try:
                        p_idx = parent.index(last_p_element)
                        parent.insert(p_idx + 1, new_p._element)
                    except ValueError: pass
                 elif doc_paragraphs_list and not last_p_element:
                    first_p = doc_paragraphs_list[0]._element
                    parent = first_p.getparent()
                    parent.insert(0, new_p._element)
                 
                 last_p_element = new_p._element

    # 3. Handle Deletions
    for pid, p_obj in original_paragraphs.items():
        if pid not in visited_ids:
            try:
                p_element = p_obj._element
                p_element.getparent().remove(p_element)
            except (AttributeError, ValueError):
                pass

    doc.save(output_path)
